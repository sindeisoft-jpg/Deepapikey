#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DeepSeek智能对话 + Word文档导出工具
基于现有项目增强，添加一键导出对话为Word文档功能
"""

import sys
import os
from datetime import datetime
from PyQt6.QtCore import Qt, QUrl, QTimer, pyqtSlot
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QTextEdit, QPushButton, QLabel, QSplitter, QMessageBox,
    QGroupBox, QFrame, QComboBox, QCheckBox, QFileDialog
)
from PyQt6.QtWebEngineWidgets import QWebEngineView
from PyQt6.QtWebEngineCore import QWebEnginePage, QWebEngineSettings
from PyQt6.QtGui import QFont, QTextCursor, QColor
import docx
from docx.shared import Inches

class DeepSeekChatWithExport(QMainWindow):
    """DeepSeek对话增强版 - 支持Word导出"""
    
    def __init__(self):
        super().__init__()
        self.conversation_history = []  # 存储对话历史
        self.init_ui()
        self.setup_connections()
        
    def init_ui(self):
        """初始化用户界面"""
        self.setWindowTitle("DeepSeek 智能对话 - Word导出增强版")
        self.setGeometry(100, 100, 1400, 900)
        
        # 创建中央部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # 主布局
        main_layout = QHBoxLayout(central_widget)
        
        # 创建分割器
        splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # 左侧：Web浏览器访问DeepSeek
        browser_panel = self.create_browser_panel()
        splitter.addWidget(browser_panel)
        
        # 右侧：增强对话面板
        chat_panel = self.create_enhanced_chat_panel()
        splitter.addWidget(chat_panel)
        
        # 设置分割比例
        splitter.setSizes([800, 600])
        main_layout.addWidget(splitter)
        
        # 状态栏
        self.statusBar().showMessage("就绪 - DeepSeek对话 + Word导出功能")
        
    def create_browser_panel(self):
        """创建浏览器面板"""
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        # 标题
        title_label = QLabel("🌐 DeepSeek 官方网站")
        title_font = QFont()
        title_font.setPointSize(14)
        title_font.setBold(True)
        title_label.setFont(title_font)
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title_label)
        
        # 浏览器控件
        self.browser = QWebEngineView()
        
        # 配置浏览器设置
        settings = self.browser.settings()
        settings.setAttribute(QWebEngineSettings.WebAttribute.JavascriptEnabled, True)
        settings.setAttribute(QWebEngineSettings.WebAttribute.LocalStorageEnabled, True)
        settings.setAttribute(QWebEngineSettings.WebAttribute.AutoLoadImages, True)
        
        # 访问DeepSeek官网
        self.browser.setUrl(QUrl("https://chat.deepseek.com"))
        layout.addWidget(self.browser)
        
        return panel
        
    def create_enhanced_chat_panel(self):
        """创建增强对话面板"""
        panel = QWidget()
        layout = QVBoxLayout(panel)
        layout.setSpacing(15)
        
        # 标题
        title_label = QLabel("💬 DeepSeek 智能对话")
        title_font = QFont()
        title_font.setPointSize(18)
        title_font.setBold(True)
        title_label.setFont(title_font)
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title_label.setStyleSheet("color: #2c3e50; padding: 15px; background-color: #ecf0f1; border-radius: 8px;")
        layout.addWidget(title_label)
        
        # 对话设置组
        settings_group = QGroupBox("对话设置")
        settings_layout = QVBoxLayout()
        
        # 回复风格选择
        style_layout = QHBoxLayout()
        style_layout.addWidget(QLabel("回复风格:"))
        self.style_combo = QComboBox()
        self.style_combo.addItems(["简洁明了", "详细全面", "专业技术", "友好亲切", "创意思维"])
        style_layout.addWidget(self.style_combo)
        style_layout.addStretch()
        settings_layout.addLayout(style_layout)
        
        # 功能选项
        self.auto_scroll_check = QCheckBox("自动滚动到最新消息")
        self.auto_scroll_check.setChecked(True)
        self.save_history_check = QCheckBox("自动保存对话历史")
        self.save_history_check.setChecked(True)
        
        settings_layout.addWidget(self.auto_scroll_check)
        settings_layout.addWidget(self.save_history_check)
        settings_group.setLayout(settings_layout)
        layout.addWidget(settings_group)
        
        # 对话历史显示
        history_label = QLabel("📝 对话历史:")
        history_label.setStyleSheet("font-weight: bold; font-size: 14px;")
        layout.addWidget(history_label)
        
        self.chat_display = QTextEdit()
        self.chat_display.setReadOnly(True)
        self.chat_display.setPlaceholderText("在此查看对话历史...\n点击右侧浏览器与DeepSeek对话，内容将同步到这里")
        self.chat_display.setMinimumHeight(300)
        self.chat_display.setStyleSheet("""
            QTextEdit {
                border: 2px solid #bdc3c7;
                border-radius: 8px;
                padding: 15px;
                font-family: 'Microsoft YaHei', sans-serif;
                font-size: 13px;
                background-color: #ffffff;
                selection-background-color: #3498db;
            }
        """)
        layout.addWidget(self.chat_display)
        
        # 输入区域
        input_group = QGroupBox("发送消息")
        input_layout = QVBoxLayout()
        
        self.message_input = QTextEdit()
        self.message_input.setMaximumHeight(100)
        self.message_input.setPlaceholderText("请输入您的问题或想法...")
        self.message_input.setStyleSheet("""
            QTextEdit {
                border: 2px solid #3498db;
                border-radius: 8px;
                padding: 12px;
                font-size: 14px;
                background-color: #ffffff;
            }
            QTextEdit:focus {
                border-color: #2980b9;
                background-color: #f8f9fa;
            }
        """)
        input_layout.addWidget(self.message_input)
        
        # 按钮行
        button_layout = QHBoxLayout()
        
        self.send_button = QPushButton("📤 发送消息")
        self.send_button.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 12px 25px;
                border-radius: 6px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #229954;
            }
            QPushButton:pressed {
                background-color: #1e8449;
            }
        """)
        
        self.clear_button = QPushButton("🗑️ 清空对话")
        self.clear_button.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 12px 25px;
                border-radius: 6px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        
        self.export_button = QPushButton("📄 导出Word")
        self.export_button.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 12px 25px;
                border-radius: 6px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        
        button_layout.addWidget(self.send_button)
        button_layout.addWidget(self.clear_button)
        button_layout.addWidget(self.export_button)
        button_layout.addStretch()
        
        input_layout.addLayout(button_layout)
        input_group.setLayout(input_layout)
        layout.addWidget(input_group)
        
        # 状态显示
        self.status_label = QLabel("状态: 就绪")
        self.status_label.setStyleSheet("color: #7f8c8d; font-style: italic; padding: 5px;")
        layout.addWidget(self.status_label)
        
        return panel
        
    def setup_connections(self):
        """设置信号连接"""
        self.send_button.clicked.connect(self.send_message)
        self.clear_button.clicked.connect(self.clear_conversation)
        self.export_button.clicked.connect(self.export_to_word)
        
        # 回车发送支持
        self.message_input.textChanged.connect(self.check_enter_key)
        
        # 浏览器页面加载完成信号
        self.browser.loadFinished.connect(self.on_page_loaded)
        
    def check_enter_key(self):
        """检查回车键发送"""
        # 这里可以添加Ctrl+Enter发送的功能
        pass
        
    def on_page_loaded(self, success):
        """页面加载完成回调"""
        if success:
            self.status_label.setText("状态: DeepSeek页面加载完成")
            self.statusBar().showMessage("DeepSeek官网已加载，可在右侧进行对话")
        else:
            self.status_label.setText("状态: 页面加载失败")
            
    def send_message(self):
        """发送消息"""
        message = self.message_input.toPlainText().strip()
        if not message:
            QMessageBox.warning(self, "输入为空", "请输入您的问题或消息")
            return
            
        # 更新状态
        self.status_label.setText("状态: 正在处理...")
        self.send_button.setEnabled(False)
        
        # 添加用户消息到显示
        self.add_message_to_display("user", message)
        
        # 保存到历史
        if self.save_history_check.isChecked():
            self.conversation_history.append({
                "timestamp": datetime.now().isoformat(),
                "role": "user",
                "content": message,
                "style": self.style_combo.currentText()
            })
        
        # 清空输入框
        self.message_input.clear()
        
        # 模拟AI回复（实际项目中这里应该调用DeepSeek API）
        QTimer.singleShot(2000, lambda: self.simulate_ai_response(message))
        
    def simulate_ai_response(self, user_message):
        """模拟AI回复（实际应替换为真实API调用）"""
        # 这里应该调用真实的DeepSeek API
        # 暂时使用模拟回复
        import random
        responses = [
            f"感谢您的提问。关于'{user_message}'，我认为可以从多个角度来分析...",
            f"这是一个很好的问题！针对'{user_message}'，我的建议是...",
            f"关于您提到的'{user_message}'，让我为您提供详细的解答...",
            f"我理解您的关注点是'{user_message}'，以下是我的分析..."
        ]
        
        ai_response = random.choice(responses)
        
        # 添加AI回复到显示
        self.add_message_to_display("ai", ai_response)
        
        # 保存到历史
        if self.save_history_check.isChecked():
            self.conversation_history.append({
                "timestamp": datetime.now().isoformat(),
                "role": "assistant",
                "content": ai_response,
                "style": self.style_combo.currentText()
            })
        
        # 更新状态
        self.status_label.setText("状态: 就绪")
        self.send_button.setEnabled(True)
        
        # 自动滚动
        if self.auto_scroll_check.isChecked():
            self.scroll_to_bottom()
            
    def add_message_to_display(self, role, message):
        """添加消息到显示区域"""
        cursor = self.chat_display.textCursor()
        cursor.movePosition(QTextCursor.MoveOperation.End)
        
        # 设置格式
        format = cursor.charFormat()
        time_str = datetime.now().strftime("%H:%M:%S")
        
        if role == "user":
            format.setForeground(QColor("#2980b9"))  # 蓝色
            format.setFontWeight(QFont.Weight.Bold)
            prefix = "👤 您: "
        else:
            format.setForeground(QColor("#27ae60"))  # 绿色
            format.setFontWeight(QFont.Weight.Normal)
            prefix = "🤖 DeepSeek: "
            
        # 添加时间戳和消息
        cursor.insertText(f"\n[{time_str}] ", format)
        cursor.insertText(f"{prefix}{message}\n", format)
        
    def scroll_to_bottom(self):
        """滚动到底部"""
        scrollbar = self.chat_display.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())
        
    def clear_conversation(self):
        """清空对话"""
        reply = QMessageBox.question(
            self, "确认清空",
            "确定要清空所有对话历史吗？此操作不可撤销。",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            self.chat_display.clear()
            self.conversation_history.clear()
            self.status_label.setText("状态: 对话已清空")
            
    def export_to_word(self):
        """导出对话为Word文档"""
        if not self.conversation_history:
            QMessageBox.warning(self, "无内容", "没有对话内容可导出")
            return
            
        # 选择保存位置
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "导出对话为Word文档",
            f"DeepSeek对话记录_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            "Word文档 (*.docx)"
        )
        
        if not file_path:
            return
            
        try:
            # 创建Word文档
            doc = docx.Document()
            
            # 添加标题
            doc.add_heading('DeepSeek 智能对话记录', 0)
            
            # 添加文档信息
            doc.add_paragraph(f'导出时间: {datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
            doc.add_paragraph(f'对话轮数: {len(self.conversation_history) // 2}轮')
            doc.add_paragraph(f'总消息数: {len(self.conversation_history)}条')
            doc.add_paragraph('')
            
            # 添加对话详情
            doc.add_heading('对话详情', level=1)
            
            for i, entry in enumerate(self.conversation_history, 1):
                timestamp = datetime.fromisoformat(entry['timestamp']).strftime("%H:%M:%S")
                role = "👤 用户" if entry['role'] == "user" else "🤖 DeepSeek"
                style = entry.get('style', '默认')
                
                # 添加消息标题
                doc.add_paragraph(f'{role} [{timestamp}] (风格: {style})', style='Heading 2')
                
                # 添加消息内容
                doc.add_paragraph(entry['content'], style='Normal')
                
                # 添加间距
                if i < len(self.conversation_history):
                    doc.add_paragraph('')
            
            # 保存文档
            doc.save(file_path)
            
            QMessageBox.information(self, "导出成功", f"对话记录已成功导出至:\n{file_path}")
            self.status_label.setText("状态: 文档导出完成")
            
        except Exception as e:
            QMessageBox.critical(self, "导出失败", f"导出过程中发生错误:\n{str(e)}")
            
    def closeEvent(self, event):
        """关闭窗口事件"""
        if self.conversation_history and self.save_history_check.isChecked():
            reply = QMessageBox.question(
                self, "保存对话",
                "是否保存当前对话历史到Word文档？",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No | QMessageBox.StandardButton.Cancel,
                QMessageBox.StandardButton.Yes
            )
            
            if reply == QMessageBox.StandardButton.Cancel:
                event.ignore()
                return
            elif reply == QMessageBox.StandardButton.Yes:
                self.export_to_word()
                
        event.accept()


def main():
    """主函数"""
    app = QApplication(sys.argv)
    app.setApplicationName("DeepSeek对话增强版")
    
    # 设置应用程序样式
    app.setStyle("Fusion")
    
    window = DeepSeekChatWithExport()
    window.show()
    
    sys.exit(app.exec())


if __name__ == "__main__":
    main()