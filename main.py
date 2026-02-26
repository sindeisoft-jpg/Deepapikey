#!/usr/bin/env python3
"""
DeepSeek Qt浏览器应用
一个内嵌Web浏览器，打开deepseek官网，并提供对话界面的Qt应用
"""

import sys
import os
from queue import Queue
from PyQt6.QtCore import Qt, QUrl, QTimer, pyqtSlot
try:
    from pynput.keyboard import Controller as KeyController, Key
    _HAS_PYNPUT = True
except ImportError:
    _HAS_PYNPUT = False
try:
    from api_server import start_api_server
    _HAS_API_SERVER = True
except ImportError:
    _HAS_API_SERVER = False
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QTextEdit, QPushButton, QLabel, QSplitter, QMessageBox, QComboBox
)
from PyQt6.QtWebEngineWidgets import QWebEngineView
from PyQt6.QtWebEngineCore import QWebEnginePage, QWebEngineSettings
from PyQt6.QtGui import QFont, QIcon


class DeepSeekBrowser(QMainWindow):
    """主窗口类，包含浏览器和对话界面"""
    
    def __init__(self):
        super().__init__()
        self._reply_stream_timer = None
        self._stream_history = ""
        self._last_reply_text = ""
        self._last_sent_message = ""  # 本次发送的用户内容，用于避免把用户消息当回复
        self._stream_unchanged_count = 0
        self._stream_poll_count = 0
        self._api_request_id = None
        self._api_response_event = None
        self._api_response_dict = None
        self._api_request_queue = None
        self._api_final_fetch_safety_timer = None  # 防止 runJavaScript 回调不触发导致第二次请求无法取到
        self._api_poll_timer = None
        self.init_ui()
        self.setup_connections()
        
    def init_ui(self):
        """初始化用户界面"""
        self.setWindowTitle("DeepSeek Qt浏览器")
        self.setGeometry(100, 100, 1200, 800)
        
        # 创建中央部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # 主布局
        main_layout = QHBoxLayout(central_widget)
        
        # 创建分割器，左侧是浏览器，右侧是对话界面
        splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # 左侧：Web浏览器面板
        browser_panel = self.create_browser_panel()
        splitter.addWidget(browser_panel)
        
        # 右侧：对话面板
        chat_panel = QWidget()
        chat_layout = QVBoxLayout(chat_panel)
        
        # 标题
        title_label = QLabel("DeepSeek 对话界面")
        title_font = QFont()
        title_font.setPointSize(16)
        title_font.setBold(True)
        title_label.setFont(title_font)
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        chat_layout.addWidget(title_label)
        
        # 输入框标签
        input_label = QLabel("输入您的问题：")
        chat_layout.addWidget(input_label)
        
        # 输入框
        self.input_text = QTextEdit()
        self.input_text.setPlaceholderText("在这里输入您的问题...")
        self.input_text.setMaximumHeight(150)
        self.input_text.setStyleSheet("""
            QTextEdit {
                border: 1px solid #ccc;
                border-radius: 5px;
                padding: 10px;
                font-size: 14px;
                color: #333333;  /* 确保文字颜色可见 */
                background-color: #ffffff;
            }
            QTextEdit:focus {
                border-color: #4CAF50;
                background-color: #f8f9fa;
            }
        """)
        chat_layout.addWidget(self.input_text)
        
        # 发送按钮
        self.send_button = QPushButton("发送")
        self.send_button.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 10px;
                font-size: 14px;
                font-weight: bold;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
            QPushButton:pressed {
                background-color: #3d8b40;
            }
        """)
        chat_layout.addWidget(self.send_button)
        
        # 返回框标签
        output_label = QLabel("DeepSeek 回复：")
        chat_layout.addWidget(output_label)
        
        # 返回框
        self.output_text = QTextEdit()
        self.output_text.setReadOnly(True)
        self.output_text.setPlaceholderText("DeepSeek的回复将显示在这里...")
        chat_layout.addWidget(self.output_text)
        
        # 控制按钮区域
        control_layout = QHBoxLayout()
        
        self.clear_button = QPushButton("清空")
        self.clear_button.setStyleSheet("""
            QPushButton {
                background-color: #f44336;
                color: white;
                border: none;
                padding: 8px;
                font-size: 12px;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #d32f2f;
            }
        """)
        
        self.export_button = QPushButton("📄 导出Word")
        self.export_button.setStyleSheet("""
            QPushButton {
                background-color: #FF9800;
                color: white;
                border: none;
                padding: 8px 15px;
                font-size: 12px;
                font-weight: bold;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #F57C00;
            }
            QPushButton:pressed {
                background-color: #E65100;
            }
        """)
        
        self.refresh_button = QPushButton("刷新页面")
        self.refresh_button.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                border: none;
                padding: 8px;
                font-size: 12px;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #1976D2;
            }
        """)
        
        control_layout.addWidget(self.clear_button)
        control_layout.addWidget(self.export_button)
        control_layout.addWidget(self.refresh_button)
        control_layout.addStretch()
        
        chat_layout.addLayout(control_layout)
        
        # 添加右侧面板到分割器
        splitter.addWidget(chat_panel)
        
        # 设置分割器初始比例（70%浏览器，30%对话面板）
        splitter.setSizes([840, 360])
        
        # 将分割器添加到主布局
        main_layout.addWidget(splitter)
        
        # 状态栏
        self.statusBar().showMessage("就绪 - 已加载DeepSeek官网")
        
    def create_browser_panel(self):
        """创建浏览器面板，包含URL导航功能"""
        panel = QWidget()
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(0, 0, 0, 0)
        
        # 浏览器控制栏
        control_layout = QHBoxLayout()
        
        self.back_button = QPushButton("←")
        self.forward_button = QPushButton("→")
        self.refresh_button = QPushButton("刷新")
        self.home_button = QPushButton("首页")
        
        # URL地址栏
        self.url_bar = QComboBox()
        self.url_bar.setEditable(True)
        self.url_bar.addItem("https://chat.deepseek.com")
        self.url_bar.addItem("https://www.deepseek.com")
        self.url_bar.addItem("https://www.deepseek.com/zh")
        
        # 设置地址栏样式
        self.url_bar.setStyleSheet("""
            QComboBox {
                color: #000000;
                background-color: #ffffff;
                border: 1px solid #ccc;
                border-radius: 3px;
                padding: 5px;
            }
            QComboBox QAbstractItemView {
                color: #000000;
                background-color: #ffffff;
            }
        """)
        
        self.go_button = QPushButton("前往")
        
        # 设置按钮样式
        for btn in [self.back_button, self.forward_button, self.refresh_button, 
                   self.home_button, self.go_button]:
            btn.setMaximumWidth(60)
            btn.setStyleSheet("""
                QPushButton {
                    padding: 5px;
                    border: 1px solid #ccc;
                    border-radius: 3px;
                    background-color: #f0f0f0;
                }
                QPushButton:hover {
                    background-color: #e0e0e0;
                }
                QPushButton:pressed {
                    background-color: #d0d0d0;
                }
            """)
        
        # 添加控件到控制栏
        control_layout.addWidget(self.back_button)
        control_layout.addWidget(self.forward_button)
        control_layout.addWidget(self.refresh_button)
        control_layout.addWidget(self.home_button)
        control_layout.addWidget(self.url_bar)
        control_layout.addWidget(self.go_button)
        
        layout.addLayout(control_layout)
        
        # Web浏览器视图
        self.browser = QWebEngineView()
        self.browser.settings().setAttribute(QWebEngineSettings.WebAttribute.JavascriptEnabled, True)
        self.browser.setUrl(QUrl("https://chat.deepseek.com"))
        layout.addWidget(self.browser)
        
        return panel
        
    def setup_connections(self):
        """设置信号和槽的连接"""
        self.send_button.clicked.connect(self.send_message)
        self.clear_button.clicked.connect(self.clear_output)
        self.export_button.clicked.connect(self.export_to_word)
        self.refresh_button.clicked.connect(self.refresh_browser)
        
        # 浏览器导航连接
        self.back_button.clicked.connect(self.browser.back)
        self.forward_button.clicked.connect(self.browser.forward)
        self.refresh_button.clicked.connect(self.browser.reload)
        self.home_button.clicked.connect(self.go_home)
        self.go_button.clicked.connect(self.navigate_to_url)
        self.url_bar.lineEdit().returnPressed.connect(self.navigate_to_url)
        
        # 浏览器加载状态变化
        self.browser.loadStarted.connect(self.on_load_started)
        self.browser.loadFinished.connect(self.on_load_finished)
        self.browser.urlChanged.connect(self.on_url_changed)
        
    def on_load_started(self):
        """浏览器开始加载页面"""
        self.statusBar().showMessage("正在加载页面...")
        
    def on_load_finished(self, success):
        """浏览器完成加载页面"""
        if success:
            self.statusBar().showMessage("页面加载完成")
        else:
            self.statusBar().showMessage("页面加载失败")
            
    def on_url_changed(self, url):
        """URL变化时更新地址栏"""
        current_url = url.toString()
        if current_url not in [self.url_bar.itemText(i) for i in range(self.url_bar.count())]:
            self.url_bar.addItem(current_url)
        self.url_bar.setCurrentText(current_url)
        
    def go_home(self):
        """返回首页"""
        self.browser.setUrl(QUrl("https://chat.deepseek.com"))
        
    def navigate_to_url(self):
        """导航到指定URL"""
        url_text = self.url_bar.currentText().strip()
        if not url_text:
            QMessageBox.warning(self, "URL为空", "请输入有效的URL地址")
            return
            
        # URL格式处理
        if not url_text.startswith(('http://', 'https://')):
            if '.' in url_text and not url_text.startswith('www.'):
                url_text = 'https://' + url_text
            elif url_text.startswith('www.'):
                url_text = 'https://' + url_text
            else:
                url_text = 'https://chat.deepseek.com'
        
        try:
            # 验证URL格式
            from urllib.parse import urlparse
            parsed = urlparse(url_text)
            if not parsed.scheme or not parsed.netloc:
                QMessageBox.warning(self, "无效URL", "请输入有效的URL地址")
                return
        except Exception as e:
            QMessageBox.warning(self, "URL错误", f"URL格式错误: {str(e)}")
            return
            
        self.browser.setUrl(QUrl(url_text))
        self.statusBar().showMessage(f"正在导航到: {url_text}")
            
    def _escape_for_js(self, text):
        """将文本转义后安全放入 JavaScript 单引号字符串中"""
        if not text:
            return ""
        return (text.replace("\\", "\\\\")
                    .replace("'", "\\'")
                    .replace("\r", "\\r")
                    .replace("\n", "\\n"))
    
    def _build_inject_script(self, message: str) -> str:
        """生成将 message 注入网页并触发发送的 JS。供 UI 与 API 共用。"""
        message_escaped = self._escape_for_js(message)
        return f"""
        (function() {{
            var msg = '{message_escaped}';
            setTimeout(function() {{
            var selectors = [
                'textarea[placeholder*="DeepSeek"]',
                'textarea[placeholder*="发送消息"]',
                'textarea[placeholder*="输入"]',
                'textarea[placeholder*="message"]',
                'textarea[placeholder*="说点什么"]',
                'textarea',
                'input[type="text"]',
                '[contenteditable="true"]',
                '[role="textbox"]',
                '.ProseMirror'
            ];
            var target = null;
            for (var i = 0; i < selectors.length; i++) {{
                var list = document.querySelectorAll(selectors[i]);
                for (var j = 0; j < list.length; j++) {{
                    var el = list[j];
                    if (el.offsetWidth > 0 && el.offsetHeight > 0) {{
                        target = el;
                        break;
                    }}
                }}
                if (target) break;
            }}
            if (!target) {{ return false; }}
            target.focus();
            if (target.tagName === 'TEXTAREA' || target.tagName === 'INPUT') {{
                var proto = target.tagName === 'TEXTAREA' ? HTMLTextAreaElement.prototype : HTMLInputElement.prototype;
                var desc = Object.getOwnPropertyDescriptor(proto, 'value');
                if (desc && desc.set) {{
                    desc.set.call(target, msg);
                }} else {{
                    target.value = msg;
                }}
                target.dispatchEvent(new InputEvent('input', {{ data: msg, inputType: 'insertText', bubbles: true }}));
                target.dispatchEvent(new Event('input', {{ bubbles: true }}));
                target.dispatchEvent(new Event('change', {{ bubbles: true }}));
            }} else if (target.contentEditable === 'true' || target.getAttribute('role') === 'textbox') {{
                target.innerText = msg;
                target.textContent = msg;
                target.dispatchEvent(new InputEvent('input', {{ data: msg, inputType: 'insertText', bubbles: true }}));
                target.dispatchEvent(new Event('input', {{ bubbles: true }}));
            }} else {{
                target.value = msg;
                target.innerText = msg;
                target.textContent = msg;
                target.dispatchEvent(new Event('input', {{ bubbles: true }}));
                target.dispatchEvent(new Event('change', {{ bubbles: true }}));
            }}
            setTimeout(function() {{
                // 简化方案：直接模拟回车键
                console.log('直接模拟回车键发送...');
                
                // 确保输入框有焦点
                target.focus();
                
                // 延迟一小段时间确保焦点稳定
                setTimeout(function() {{
                    // 创建完整的回车键事件序列
                    var events = [
                        new KeyboardEvent('keydown', {{
                            key: 'Enter',
                            code: 'Enter',
                            keyCode: 13,
                            which: 13,
                            bubbles: true,
                            cancelable: true
                        }}),
                        new KeyboardEvent('keypress', {{
                            key: 'Enter',
                            code: 'Enter',
                            keyCode: 13,
                            which: 13,
                            bubbles: true,
                            cancelable: true
                        }}),
                        new KeyboardEvent('keyup', {{
                            key: 'Enter',
                            code: 'Enter',
                            keyCode: 13,
                            which: 13,
                            bubbles: true,
                            cancelable: true
                        }})
                    ];
                    
                    // 依次触发所有事件
                    events.forEach(function(event, index) {{
                        console.log('触发事件', index + 1, ':', event.type);
                        target.dispatchEvent(event);
                    }});
                    
                    console.log('✅ 回车键模拟完成');
                }}, 100);
                
            }}, 300);
            }}, 150);
            return true;
        }})();
        """

    def send_message(self):
        """发送消息：将右侧输入提交到左侧 DeepSeek 网页的输入框并触发发送"""
        message = self.input_text.toPlainText().strip()
        if not message:
            QMessageBox.warning(self, "输入为空", "请输入您的问题")
            return
        self.output_text.append(f"您: {message}")
        self.output_text.append("")
        self.output_text.verticalScrollBar().setValue(
            self.output_text.verticalScrollBar().maximum()
        )
        self.input_text.clear()
        self.statusBar().showMessage("正在发送到网页...")
        self._last_sent_message = message
        self.browser.page().runJavaScript(
            self._build_inject_script(message), self._on_web_send_done
        )

    def _inject_and_send(self, message: str):
        """仅注入并发送到网页（供 API 调用，不更新右侧输入/输出）。"""
        self.statusBar().showMessage("API 请求处理中…")
        self._stream_history = self.output_text.toPlainText()
        self._last_reply_text = ""
        self._last_sent_message = message
        self._stream_unchanged_count = 0
        self._stream_poll_count = 0
        if self._reply_stream_timer is not None:
            self._reply_stream_timer.stop()
        self.browser.page().runJavaScript(
            self._build_inject_script(message), self._on_web_send_done
        )
    
    def _on_web_send_done(self, success):
        """网页注入/点击完成后的回调"""
        if success:
            self.statusBar().showMessage("已发送到网页，等待回复…")
            self._stream_history = self.output_text.toPlainText()
            self._last_reply_text = ""
            self._stream_unchanged_count = 0
            self._stream_poll_count = 0
            if self._reply_stream_timer is not None:
                self._reply_stream_timer.stop()
            QTimer.singleShot(500, self._simulate_enter_key)
            QTimer.singleShot(1500, self._start_reply_stream)
            
            # 添加调试信息
            debug_script = '''
            console.log("=== 调试信息 ===");
            console.log("页面标题:", document.title);
            console.log("当前URL:", window.location.href);
            console.log("页面状态:", document.readyState);
            
            // 检查输入框状态
            var inputElements = document.querySelectorAll("textarea, input[type='text'], [contenteditable='true']");
            console.log("找到输入元素数量:", inputElements.length);
            
            for(var i = 0; i < inputElements.length; i++) {
                var el = inputElements[i];
                console.log("元素" + i + ":", el.tagName, el.className, "可见:", el.offsetWidth > 0 && el.offsetHeight > 0, "值:", el.value || el.innerText);
            }
            
            // 检查按钮
            var buttons = document.querySelectorAll("button");
            console.log("找到按钮数量:", buttons.length);
            
            return {
                pageTitle: document.title,
                pageUrl: window.location.href,
                inputCount: inputElements.length,
                buttonCount: buttons.length
            };
            '''
            
            def debug_callback(result):
                print(f"调试结果: {result}")
                
            self.browser.page().runJavaScript(debug_script, debug_callback)
            
        else:
            self.statusBar().showMessage("未能找到网页输入框，请确认左侧已打开 DeepSeek 聊天页")
    
    def _simulate_enter_key(self):
        """用 pynput 模拟按下 Enter（系统级，页面会视为真实按键）"""
        self.activateWindow()
        self.raise_()
        self.browser.setFocus()
        if _HAS_PYNPUT:
            try:
                ctrl = KeyController()
                ctrl.press(Key.enter)
                ctrl.release(Key.enter)
            except Exception:
                pass

    def _start_reply_stream(self):
        """开始轮询网页中的回复，以流式方式更新到右侧 - 优化版"""
        self.statusBar().showMessage("正在实时获取网页回复...")
        self._stream_unchanged_count = 0
        self._stream_poll_count = 0
        
        # 初始化流式显示相关变量
        self._current_displayed_text = ""  # 当前已显示的文本
        self._last_code_block = ""         # 上次检测到的代码块
        self._code_blocks_found = []       # 已发现的代码块列表
        
        if self._reply_stream_timer is None:
            self._reply_stream_timer = QTimer(self)
            self._reply_stream_timer.timeout.connect(self._poll_reply)
        
        # 更频繁的轮询以获得更好的实时性（200ms）
        self._reply_stream_timer.start(200)
        
        # 添加实时显示指示器
        self._add_stream_indicator()

    def _add_stream_indicator(self):
        """添加实时流式显示指示器"""
        # 在状态栏添加流式指示器
        if not hasattr(self, '_stream_indicator'):
            self._stream_indicator = QLabel("● 实时流式传输中")
            self._stream_indicator.setStyleSheet("""
                QLabel {
                    color: #4CAF50;
                    font-weight: bold;
                    animation: blink 1s infinite;
                }
                @keyframes blink {
                    0%, 50% { opacity: 1; }
                    51%, 100% { opacity: 0.5; }
                }
            """)
            self.statusBar().addPermanentWidget(self._stream_indicator)
    
    def _remove_stream_indicator(self):
        """移除流式显示指示器"""
        if hasattr(self, '_stream_indicator'):
            self._stream_indicator.setParent(None)
            delattr(self, '_stream_indicator')
    
    def _enhanced_stream_update(self, new_content):
        """增强的流式更新逻辑，专门优化代码显示"""
        if not new_content:
            return
            
        # 检测新的代码块
        current_code_blocks = self._extract_code_blocks(new_content)
        
        # 如果发现了新的代码块或代码块有更新
        if self._detect_code_changes(current_code_blocks):
            self._display_code_progress(current_code_blocks)
            
        # 检测普通文本的变化
        if len(new_content) > len(self._current_displayed_text):
            # 只显示新增的部分
            new_text = new_content[len(self._current_displayed_text):]
            if new_text.strip():
                self._append_to_output("assistant", new_text, is_incremental=True)
                self._current_displayed_text = new_content
                
        # 滚动到底部保持最新内容可见
        self._scroll_to_latest()
    
    def _extract_code_blocks(self, content):
        """提取内容中的代码块"""
        import re
        # 匹配代码块的正则表达式
        code_pattern = r'```(?:\w+)?\s*\n([\s\S]*?)\n```|`([^`]+)`'
        matches = re.findall(code_pattern, content)
        
        code_blocks = []
        for match in matches:
            if match[0]:  # 三个反引号的代码块
                code_blocks.append(match[0].strip())
            elif match[1]:  # 单个反引号的代码
                code_blocks.append(match[1].strip())
                
        return code_blocks
    
    def _detect_code_changes(self, current_blocks):
        """检测代码块是否有变化"""
        if not current_blocks:
            return False
            
        # 检查是否有新的代码块
        if len(current_blocks) > len(self._code_blocks_found):
            return True
            
        # 检查现有代码块是否有更新
        for i, (current, previous) in enumerate(zip(current_blocks, self._code_blocks_found)):
            if current != previous and len(current) > len(previous):
                return True
                
        return False
    
    def _display_code_progress(self, code_blocks):
        """显示代码编写进度"""
        for i, code_block in enumerate(code_blocks):
            if i >= len(self._code_blocks_found):
                # 新的代码块
                self._append_to_output("code_start", f"开始编写第{i+1}个代码块...")
                self._code_blocks_found.append("")
                
            if code_block != self._code_blocks_found[i]:
                # 代码块有更新
                new_content = code_block[len(self._code_blocks_found[i]):]
                if new_content:
                    self._append_to_output("code_progress", new_content, is_code=True)
                    self._code_blocks_found[i] = code_block
                    
                    # 更新状态栏显示进度
                    progress = len(code_block) / max(1, len(code_block) + 50) * 100  # 估算进度
                    self.statusBar().showMessage(f"代码编写中... ({progress:.0f}%)")
    
    def _append_to_output(self, role, content, is_incremental=False, is_code=False):
        """向输出区域追加内容"""
        cursor = self.output_text.textCursor()
        cursor.movePosition(cursor.MoveOperation.End)
        
        # 设置不同的格式
        format_obj = cursor.charFormat()
        
        if role == "user":
            format_obj.setForeground(QColor("#2196F3"))
            format_obj.setFontWeight(QFont.Weight.Bold)
            prefix = "您: "
        elif role == "assistant":
            format_obj.setForeground(QColor("#4CAF50"))
            prefix = "DeepSeek: "
        elif role == "code_start":
            format_obj.setForeground(QColor("#FF9800"))
            format_obj.setFontWeight(QFont.Weight.Bold)
            prefix = "🔧 "
        elif role == "code_progress":
            format_obj.setForeground(QColor("#9C27B0"))
            format_obj.setFontFamily("Monaco")
            format_obj.setFontPointSize(12)
            prefix = "```\n"
            content += "\n```\n"
        else:
            prefix = ""
            
        # 添加时间戳
        from datetime import datetime
        time_str = datetime.now().strftime("%H:%M:%S")
        cursor.insertText(f"[{time_str}] ", format_obj)
        
        # 添加内容
        if prefix:
            cursor.insertText(prefix, format_obj)
        cursor.insertText(content, format_obj)
        cursor.insertText("\n\n", format_obj)
        
        # 如果是增量更新且不是代码，则添加分隔符
        if is_incremental and not is_code:
            cursor.insertText("─" * 30 + "\n", format_obj)
    
    def _scroll_to_latest(self):
        """滚动到最新内容"""
        scrollbar = self.output_text.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())

    def _poll_reply(self):
        """从网页抓取当前「最后一条」助手回复，避免第二次及以后取到第一条数据"""
        self.browser.page().runJavaScript(self._get_reply_script(), self._on_reply_chunk)

    def _on_reply_chunk(self, reply_str):
        """收到网页返回的回复片段，流式更新右侧显示；API 请求在稳定后做一次最终抓取再返回。"""
        self._stream_poll_count += 1
        if self._stream_poll_count > 200:
            self._stop_reply_stream()
            self._flush_api_response_if_any()
            return
        if reply_str is None:
            reply_str = ""
        if not isinstance(reply_str, str):
            reply_str = str(reply_str) if reply_str else ""
        reply_str = reply_str.strip()
        if reply_str and reply_str.strip() == (self._last_sent_message or "").strip():
            return
        # 防止 DOM 短暂切到其它节点导致内容突然变短（断断续续）
        if self._last_reply_text and len(reply_str) < len(self._last_reply_text) - 100:
            if len(reply_str) < max(100, int(len(self._last_reply_text) * 0.8)):
                return
        if reply_str == self._last_reply_text:
            self._stream_unchanged_count += 1
            if self._stream_unchanged_count >= 8:
                # 稳定 8 次（约 4s）后再做一次最终抓取，再写入 API 响应，避免用中间状态
                if self._api_request_id and self._api_response_dict is not None:
                    self._reply_stream_timer.stop()
                    # 若 runJavaScript 回调未触发，8s 后强制写回并清空状态，避免第二次请求永远不执行
                    if self._api_final_fetch_safety_timer is not None:
                        self._api_final_fetch_safety_timer.stop()
                    self._api_final_fetch_safety_timer = QTimer(self)
                    self._api_final_fetch_safety_timer.setSingleShot(True)
                    self._api_final_fetch_safety_timer.timeout.connect(self._api_safety_flush_and_clear)
                    self._api_final_fetch_safety_timer.start(8000)
                    QTimer.singleShot(600, self._final_fetch_for_api)  # 留时间让代码块渲染完再抓
                else:
                    self._stop_reply_stream()
                    if reply_str:
                        self._stream_history = self.output_text.toPlainText() + "\n\n"
                return
            return
        self._stream_unchanged_count = 0
        self._last_reply_text = reply_str
        if self._api_request_id is None:
            display = self._stream_history + "DeepSeek: " + reply_str
            self.output_text.setPlainText(display)
            self.output_text.verticalScrollBar().setValue(
                self.output_text.verticalScrollBar().maximum()
            )

    def _flush_api_response_if_any(self):
        """超时或停止时，若有未完成的 API 请求则用当前 _last_reply_text 写回并 set event。"""
        if not self._api_request_id or self._api_response_dict is None:
            return
        self._api_response_dict[self._api_request_id] = self._last_reply_text or ""
        if self._api_response_event:
            self._api_response_event.set()
        self._api_request_id = None
        self._api_response_event = None
        self.statusBar().showMessage("API 请求已完成")

    def _final_fetch_for_api(self):
        """稳定后做一次最终抓取，用此次结果作为 API 的 content，再停止轮询并写回。"""
        self.browser.page().runJavaScript(self._get_reply_script(), self._on_final_fetch_done)

    def _get_reply_script(self):
        """从页面抓取最后一条助手回复（含代码块），优先取整条消息根节点再提取全文。"""
        return """
        (function() {
            function getText(el) {
                if (!el) return '';
                var t = el.innerText || el.textContent || '';
                return (typeof t === 'string' ? t : '').trim();
            }
            function getCodeBlockLang(node) {
                var code = node.tagName === 'CODE' ? node : node.querySelector('code');
                var el = code || node;
                var cls = (el.className || '') + ' ' + (el.getAttribute('class') || '');
                var m = cls.match(/language-(\\w+)/);
                return m ? m[1] : '';
            }
            function toMarkdownLike(el) {
                if (!el) return '';
                var out = [];
                function walk(node) {
                    if (node.nodeType === 1) {
                        var tag = (node.tagName || '').toUpperCase();
                        if (tag === 'PRE') {
                            var code = node.querySelector('code');
                            var block = (code || node).innerText || (code || node).textContent || '';
                            var lang = getCodeBlockLang(node);
                            if (block) out.push('```' + (lang || '') + '\\n' + block.trim() + '\\n```');
                            return;
                        }
                        if (tag === 'DIV' && node.querySelector && !node.querySelector('pre')) {
                            var c = (node.className || '') + ' ' + (node.getAttribute('class') || '');
                            if (/code|Code|highlight/.test(c)) {
                                var block = getText(node);
                                if (block && block.length > 2) out.push('```\\n' + block + '\\n```');
                                return;
                            }
                        }
                        for (var i = 0; i < node.childNodes.length; i++) walk(node.childNodes[i]);
                    } else if (node.nodeType === 3) {
                        var t = (node.textContent || '').trim();
                        if (t) out.push(t);
                    }
                }
                walk(el);
                return out.join('\\n\\n').trim() || getText(el);
            }
            function isWelcome(t) {
                return t.indexOf('今天有什么可以帮') >= 0 || t.indexOf('有什么可以帮') >= 0;
            }
            function inDocOrder(a, b) {
                var pos = a.compareDocumentPosition(b);
                return (pos & Node.DOCUMENT_POSITION_FOLLOWING) ? -1 : 1;
            }
            // 1) 优先：只匹配「整条消息」的根节点（一条消息一个节点），避免取到不含代码的子块
            var rootSel = [
                '[data-message-type="assistant"]',
                '[class*="assistant"][class*="message"]',
                '[class*="message"][class*="assistant"]',
                '[role="article"][class*="assistant"]',
                'article[class*="assistant"]',
                'div[class*="assistant"][class*="message"]'
            ];
            for (var s = 0; s < rootSel.length; s++) {
                try {
                    var list = document.querySelectorAll(rootSel[s]);
                    if (list.length === 0) continue;
                    var roots = [];
                    for (var i = 0; i < list.length; i++) {
                        var el = list[i];
                        var t = getText(el);
                        if (t.length < 3 || isWelcome(t)) continue;
                        roots.push(el);
                    }
                    if (roots.length > 0) {
                        roots.sort(inDocOrder);
                        var lastRoot = roots[roots.length - 1];
                        return toMarkdownLike(lastRoot) || getText(lastRoot);
                    }
                } catch (e) {}
            }
            // 2) 回退：从任意节点中找「最后一条助手消息」的根（closest 到 message 根）
            var anySel = [
                '[class*="message"]', '[class*="Message"]',
                '[class*="assistant"]', '[class*="markdown"]', '[class*="content"]', '[class*="prose"]',
                'article', '[role="article"]', '[data-message-type="assistant"]', '[class*="reply"]',
                'pre', 'code'
            ];
            var candidates = [];
            for (var i = 0; i < anySel.length; i++) {
                var els = document.querySelectorAll(anySel[i]);
                for (var j = 0; j < els.length; j++) {
                    var el = els[j];
                    var t = getText(el);
                    if (t.length < 3 || isWelcome(t)) continue;
                    var root = el.closest && (
                        el.closest('[data-message-type="assistant"]') ||
                        el.closest('[class*="message"][class*="assistant"]') ||
                        el.closest('[class*="assistant"][class*="message"]') ||
                        el.closest('article[class*="assistant"]') ||
                        el.closest('[role="article"]') ||
                        el
                    );
                    var rootText = getText(root);
                    if (root && rootText.length > t.length + 20) t = rootText;
                    candidates.push({ el: root || el, text: t, len: (rootText || t).length });
                }
            }
            if (candidates.length === 0) {
                var main = document.querySelector('main') || document.querySelector('[role="main"]') || document.body;
                var full = getText(main);
                if (full.length > 50 && !isWelcome(full)) {
                    var idx = Math.max(full.lastIndexOf('您:'), full.lastIndexOf('You:'));
                    return idx >= 0 ? full.substring(idx).trim() : full;
                }
                var all = document.querySelectorAll('div, section, article');
                for (var k = all.length - 1; k >= 0; k--) {
                    var t = getText(all[k]);
                    if (t.length > 50 && t.length < 500000 && !isWelcome(t)) return t;
                }
                return '';
            }
            candidates.sort(function(a, b) { return inDocOrder(a.el, b.el); });
            var lastInDoc = candidates[candidates.length - 1].el;
            var fullContent = toMarkdownLike(lastInDoc) || getText(lastInDoc);
            if (fullContent.length > 50) return fullContent;
            var byLen = candidates.slice().sort(function(a, b) { return (b.len || b.text.length) - (a.len || a.text.length); });
            return toMarkdownLike(byLen[0].el) || byLen[0].text;
        })();
        """

    def _api_safety_flush_and_clear(self):
        """超时兜底：若最终抓取回调未触发，强制写回当前内容并清空状态，以便下一次 API 请求能执行。"""
        if self._api_final_fetch_safety_timer is not None:
            self._api_final_fetch_safety_timer.stop()
            self._api_final_fetch_safety_timer = None
        self._flush_api_response_if_any()

    def _on_final_fetch_done(self, reply_str):
        """最终抓取回调：增强错误处理和重试机制"""
        try:
            # 停止安全定时器
            if self._api_final_fetch_safety_timer is not None:
                self._api_final_fetch_safety_timer.stop()
                self._api_final_fetch_safety_timer = None
                
            # 停止回复流
            self._stop_reply_stream()
            
            # 处理回复内容
            final = (reply_str or "").strip() if isinstance(reply_str, str) else ""
            if not final:
                final = self._last_reply_text or ""
                
            # 记录调试信息
            print(f"DEBUG: API最终回复 - 长度: {len(final)}, 内容预览: {final[:100]}")
            
            # 确保API响应字典存在
            if self._api_request_id and self._api_response_dict is not None:
                self._api_response_dict[self._api_request_id] = final
                if self._api_response_event:
                    self._api_response_event.set()
                    print(f"DEBUG: API事件已设置，请求ID: {self._api_request_id}")
                self._api_request_id = None
                self._api_response_event = None
            else:
                print("DEBUG: API响应状态异常")
                # 兜底处理
                self._flush_api_response_if_any()
                
            self.statusBar().showMessage("API 请求已完成")
            
        except Exception as e:
            print(f"DEBUG: 回调处理异常: {e}")
            # 兜底处理
            self._api_safety_flush_and_clear()

    def set_api_queues(self, request_queue: Queue, response_dict: dict):
        """设置 API 请求队列与响应字典（由 main 在启动 API 服务后调用）。"""
        self._api_request_queue = request_queue
        self._api_response_dict = response_dict  # 与 api_server 共用，主线程写入回复

    def start_api_polling(self):
        """开始轮询 API 请求队列（需先 set_api_queues）。"""
        if self._api_request_queue is None or self._api_poll_timer is not None:
            return
        self._api_poll_timer = QTimer(self)
        self._api_poll_timer.timeout.connect(self._poll_api_request)
        self._api_poll_timer.start(500)

    def _poll_api_request(self):
        """从队列取 API 请求并在主线程执行注入与发送。"""
        if self._api_request_queue is None or self._api_request_id is not None:
            return
        try:
            request_id, message, event = self._api_request_queue.get_nowait()
        except Exception:
            return
        self._api_request_id = request_id
        self._api_response_event = event
        self._inject_and_send(message)

    def clear_output(self):
        """清空输出框"""
        if self._reply_stream_timer is not None:
            self._reply_stream_timer.stop()
        self._stream_history = ""
        self._last_reply_text = ""
        self.output_text.clear()
        self.statusBar().showMessage("输出已清空")
        
    def refresh_browser(self):
        """刷新浏览器页面"""
        self.browser.reload()
        self.statusBar().showMessage("正在刷新页面...")
        
    def export_to_word(self):
        """导出对话内容为Word文档"""
        try:
            # 获取对话内容
            content = self.output_text.toPlainText().strip()
            if not content:
                QMessageBox.information(self, "无内容", "没有对话内容可以导出")
                return
            
            # 选择保存位置
            from PyQt6.QtWidgets import QFileDialog
            import docx
            from datetime import datetime
            
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "导出对话为Word文档",
                f"DeepSeek对话记录_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "Word文档 (*.docx)"
            )
            
            if not file_path:
                return
            
            # 创建Word文档
            doc = docx.Document()
            
            # 添加标题
            doc.add_heading('DeepSeek 对话记录', 0)
            
            # 添加基本信息
            doc.add_paragraph(f'导出时间: {datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
            doc.add_paragraph('')
            
            # 添加对话内容
            doc.add_heading('对话内容', level=1)
            
            # 解析并格式化对话内容
            lines = content.split('\n')
            for line in lines:
                if line.strip():
                    if line.startswith('您:'):
                        # 用户消息
                        doc.add_paragraph(line, style='Heading 2')
                    elif 'DeepSeek:' in line:
                        # AI回复
                        doc.add_paragraph(line, style='Normal')
                    else:
                        # 其他内容
                        doc.add_paragraph(line, style='Normal')
                else:
                    # 空行
                    doc.add_paragraph('')
            
            # 保存文档
            doc.save(file_path)
            
            QMessageBox.information(self, "导出成功", f"对话记录已成功导出至:\n{file_path}")
            self.statusBar().showMessage("文档导出完成")
            
        except ImportError:
            QMessageBox.critical(self, "缺少依赖", "请安装python-docx库: pip install python-docx")
        except Exception as e:
            QMessageBox.critical(self, "导出失败", f"导出过程中发生错误:\n{str(e)}")
        
    def closeEvent(self, event):
        """关闭窗口事件"""
        reply = QMessageBox.question(
            self, "确认退出",
            "确定要退出DeepSeek浏览器吗？",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            event.accept()
        else:
            event.ignore()


def main():
    """主函数"""
    app = QApplication(sys.argv)
    app.setApplicationName("DeepSeek Qt浏览器")
    app.setStyle("Fusion")

    window = DeepSeekBrowser()
    if _HAS_API_SERVER:
        try:
            request_queue = Queue()
            response_dict = {}
            port = int(os.environ.get("DEEPSEEK_API_PORT", "8765"))
            start_api_server(request_queue, response_dict, port=port)
            window.set_api_queues(request_queue, response_dict)
            window.start_api_polling()
            window.statusBar().showMessage(
                f"就绪 - 已加载 DeepSeek 官网 | Ollama 兼容 API: http://127.0.0.1:{port}/ (无 API Key)"
            )
        except Exception as e:
            window.statusBar().showMessage(f"就绪 - API 未启动: {e}")
    else:
        window.statusBar().showMessage("就绪 - 已加载 DeepSeek 官网")
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()