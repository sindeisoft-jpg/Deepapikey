#!/usr/bin/env python3
"""
DeepSeek Qt浏览器增强版
包含更多功能和更好的界面优化
"""

import sys
import os
import json
import random
from datetime import datetime
from PyQt6.QtCore import Qt, QUrl, QTimer, pyqtSlot, QSize
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QTextEdit, QPushButton, QLabel, QSplitter, QMessageBox,
    QGroupBox, QFrame, QProgressBar, QComboBox, QCheckBox
)
from PyQt6.QtWebEngineWidgets import QWebEngineView
from PyQt6.QtWebEngineCore import QWebEnginePage, QWebEngineSettings
from PyQt6.QtGui import QFont, QIcon, QColor, QTextCharFormat, QTextCursor, QPalette


class EnhancedDeepSeekBrowser(QMainWindow):
    """增强版主窗口类"""
    
    def __init__(self):
        super().__init__()
        self.conversation_history = []
        self.init_ui()
        self.setup_connections()
        self.load_settings()
        
    def init_ui(self):
        """初始化用户界面"""
        self.setWindowTitle("DeepSeek Qt浏览器 - 增强版")
        self.setGeometry(100, 100, 1400, 900)
        
        # 设置应用程序图标（如果有）
        try:
            self.setWindowIcon(QIcon("icon.png"))
        except:
            pass
        
        # 创建中央部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # 主布局
        main_layout = QHBoxLayout(central_widget)
        main_layout.setContentsMargins(10, 10, 10, 10)
        main_layout.setSpacing(10)
        
        # 创建分割器
        splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # 左侧：Web浏览器面板
        browser_panel = self.create_browser_panel()
        splitter.addWidget(browser_panel)
        
        # 右侧：对话和控制面板
        chat_panel = self.create_chat_panel()
        splitter.addWidget(chat_panel)
        
        # 设置分割器初始比例
        splitter.setSizes([900, 500])
        
        # 将分割器添加到主布局
        main_layout.addWidget(splitter)
        
        # 状态栏
        self.status_label = QLabel("就绪 - 已加载DeepSeek官网")
        self.statusBar().addPermanentWidget(self.status_label)
        
        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setMaximumWidth(200)
        self.progress_bar.setVisible(False)
        self.statusBar().addPermanentWidget(self.progress_bar)
        
    def create_browser_panel(self):
        """创建浏览器面板"""
        panel = QWidget()
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(0, 0, 0, 0)
        
        # 浏览器控制栏
        control_layout = QHBoxLayout()
        
        self.back_button = QPushButton("←")
        self.forward_button = QPushButton("→")
        self.refresh_button = QPushButton("刷新")
        self.home_button = QPushButton("首页")
        
        self.url_bar = QComboBox()
        self.url_bar.setEditable(True)
        self.url_bar.addItem("https://chat.deepseek.com")
        self.url_bar.addItem("https://www.deepseek.com")
        self.url_bar.addItem("https://www.deepseek.com/zh")
        # 设置地址栏文字颜色为黑色
        self.url_bar.setStyleSheet("""
            QComboBox {
                color: #000000;  /* 黑色文字 */
                background-color: #ffffff;
                border: 1px solid #ccc;
                border-radius: 3px;
                padding: 5px;
            }
            QComboBox QAbstractItemView {
                color: #000000;  /* 下拉列表文字颜色 */
                background-color: #ffffff;
            }
        """)
        
        self.go_button = QPushButton("前往")
        
        # 设置按钮样式
        for btn in [self.back_button, self.forward_button, self.refresh_button, 
                   self.home_button, self.go_button]:
            btn.setMaximumWidth(60)
            btn.setStyleSheet("""
                QPushButton {
                    padding: 5px;
                    border: 1px solid #ccc;
                    border-radius: 3px;
                    background-color: #f5f5f5;
                }
                QPushButton:hover {
                    background-color: #e0e0e0;
                }
            """)
        
        control_layout.addWidget(self.back_button)
        control_layout.addWidget(self.forward_button)
        control_layout.addWidget(self.refresh_button)
        control_layout.addWidget(self.home_button)
        control_layout.addWidget(QLabel("地址:"))
        control_layout.addWidget(self.url_bar, 1)
        control_layout.addWidget(self.go_button)
        
        layout.addLayout(control_layout)
        
        # Web浏览器
        self.browser = QWebEngineView()
        
        # 启用开发者工具
        self.browser.settings().setAttribute(QWebEngineSettings.WebAttribute.JavascriptEnabled, True)
        self.browser.settings().setAttribute(QWebEngineSettings.WebAttribute.LocalStorageEnabled, True)
        
        self.browser.setUrl(QUrl("https://chat.deepseek.com"))
        layout.addWidget(self.browser, 1)
        
        return panel
        
    def create_chat_panel(self):
        """创建对话面板"""
        panel = QWidget()
        layout = QVBoxLayout(panel)
        layout.setSpacing(10)
        
        # 标题
        title_label = QLabel("DeepSeek 智能对话")
        title_font = QFont()
        title_font.setPointSize(18)
        title_font.setBold(True)
        title_label.setFont(title_font)
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title_label.setStyleSheet("color: #2c3e50; padding: 10px;")
        layout.addWidget(title_label)
        
        # 分隔线
        line = QFrame()
        line.setFrameShape(QFrame.Shape.HLine)
        line.setFrameShadow(QFrame.Shadow.Sunken)
        line.setStyleSheet("color: #bdc3c7;")
        layout.addWidget(line)
        
        # 设置组
        settings_group = QGroupBox("对话设置")
        settings_layout = QVBoxLayout()
        
        # 回复风格选择
        style_layout = QHBoxLayout()
        style_layout.addWidget(QLabel("回复风格:"))
        self.style_combo = QComboBox()
        self.style_combo.addItems(["简洁", "详细", "专业", "友好", "创意"])
        style_layout.addWidget(self.style_combo)
        style_layout.addStretch()
        settings_layout.addLayout(style_layout)
        
        # 选项
        self.auto_scroll_check = QCheckBox("自动滚动到最新消息")
        self.auto_scroll_check.setChecked(True)
        self.save_history_check = QCheckBox("保存对话历史")
        self.save_history_check.setChecked(True)
        
        settings_layout.addWidget(self.auto_scroll_check)
        settings_layout.addWidget(self.save_history_check)
        settings_group.setLayout(settings_layout)
        layout.addWidget(settings_group)
        
        # 对话历史框
        history_label = QLabel("对话历史:")
        layout.addWidget(history_label)
        
        self.output_text = QTextEdit()
        self.output_text.setReadOnly(True)
        self.output_text.setPlaceholderText("对话历史将显示在这里...")
        self.output_text.setMinimumHeight(300)
        
        # 设置对话历史框样式
        self.output_text.setStyleSheet("""
            QTextEdit {
                border: 1px solid #ddd;
                border-radius: 5px;
                padding: 10px;
                background-color: #f9f9f9;
                font-family: 'Microsoft YaHei', sans-serif;
            }
        """)
        layout.addWidget(self.output_text, 1)
        
        # 输入框
        input_label = QLabel("输入您的问题:")
        layout.addWidget(input_label)
        
        self.input_text = QTextEdit()
        self.input_text.setPlaceholderText("在这里输入您的问题...（按Ctrl+Enter发送）")
        self.input_text.setMaximumHeight(120)
        self.input_text.setStyleSheet("""
            QTextEdit {
                border: 2px solid #3498db;
                border-radius: 5px;
                padding: 10px;
                font-size: 14px;
                color: #333333;  /* 确保文字颜色可见 */
                background-color: #ffffff;
            }
            QTextEdit:focus {
                border-color: #2980b9;
                background-color: #f8f9fa;
            }
        """)
        layout.addWidget(self.input_text)
        
        # 按钮区域
        button_layout = QHBoxLayout()
        
        self.send_button = QPushButton("发送 (Ctrl+Enter)")
        self.send_button.setStyleSheet("""
            QPushButton {
                background-color: #2ecc71;
                color: white;
                border: none;
                padding: 12px 24px;
                font-size: 14px;
                font-weight: bold;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #27ae60;
            }
            QPushButton:pressed {
                background-color: #229954;
            }
            QPushButton:disabled {
                background-color: #95a5a6;
            }
        """)
        
        self.clear_button = QPushButton("清空历史")
        self.clear_button.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 10px 20px;
                font-size: 13px;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        
        self.export_button = QPushButton("导出对话")
        self.export_button.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 10px 20px;
                font-size: 13px;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        
        button_layout.addWidget(self.send_button)
        button_layout.addWidget(self.clear_button)
        button_layout.addWidget(self.export_button)
        button_layout.addStretch()
        
        layout.addLayout(button_layout)
        
        # 状态提示
        self.chat_status = QLabel("就绪")
        self.chat_status.setStyleSheet("color: #7f8c8d; font-style: italic;")
        layout.addWidget(self.chat_status)
        
        return panel
        
    def setup_connections(self):
        """设置信号和槽的连接"""
        # 浏览器控制
        self.back_button.clicked.connect(self.browser.back)
        self.forward_button.clicked.connect(self.browser.forward)
        self.refresh_button.clicked.connect(self.browser.reload)
        self.home_button.clicked.connect(self.go_home)
        self.go_button.clicked.connect(self.navigate_to_url)
        # QComboBox没有returnPressed信号，使用lineEdit的returnPressed
        self.url_bar.lineEdit().returnPressed.connect(self.navigate_to_url)
        
        # 对话功能
        self.send_button.clicked.connect(self.send_message)
        self.clear_button.clicked.connect(self.clear_conversation)
        self.export_button.clicked.connect(self.export_conversation)
        
        # 浏览器信号
        self.browser.loadStarted.connect(self.on_load_started)
        self.browser.loadProgress.connect(self.on_load_progress)
        self.browser.loadFinished.connect(self.on_load_finished)
        self.browser.urlChanged.connect(self.on_url_changed)
        
        # 快捷键
        self.send_button.setShortcut("Ctrl+Return")
        
    def load_settings(self):
        """加载设置"""
        # 这里可以添加从文件加载设置的代码
        pass
        
    def on_load_started(self):
        """浏览器开始加载页面"""
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        self.status_label.setText("正在加载页面...")
        
    def on_load_progress(self, progress):
        """浏览器加载进度"""
        self.progress_bar.setValue(progress)
        
    def on_load_finished(self, success):
        """浏览器完成加载页面"""
        self.progress_bar.setVisible(False)
        if success:
            self.status_label.setText(f"页面加载完成 - {self.browser.url().toString()}")
        else:
            self.status_label.setText("页面加载失败")
            QMessageBox.warning(self, "加载失败", "无法加载页面，请检查网络连接或URL地址")
            
    def on_url_changed(self, url):
        """URL变化"""
        current_url = url.toString()
        if current_url not in [self.url_bar.itemText(i) for i in range(self.url_bar.count())]:
            self.url_bar.addItem(current_url)
        self.url_bar.setCurrentText(current_url)
        
    def go_home(self):
        """返回首页"""
        self.browser.setUrl(QUrl("https://chat.deepseek.com"))
        
    def navigate_to_url(self):
        """导航到指定URL"""
        url_text = self.url_bar.currentText().strip()
        if not url_text.startswith(('http://', 'https://')):
            url_text = 'https://' + url_text
            
        self.browser.setUrl(QUrl(url_text))
        
    def send_message(self):
        """发送消息"""
        message = self.input_text.toPlainText().strip()
        
        if not message:
            QMessageBox.warning(self, "输入为空", "请输入您的问题")
            return
            
        # 禁用发送按钮，防止重复发送
        self.send_button.setEnabled(False)
        self.chat_status.setText("正在处理您的问题...")
        
        # 添加用户消息到历史
        self.add_message_to_history("user", message)
        
        # 清空输入框
        self.input_text.clear()
        
        # 模拟处理延迟
        QTimer.singleShot(1500, lambda: self.process_ai_response(message))
        
    def process_ai_response(self, user_message):
        """处理AI回复"""
        try:
            # 获取选择的回复风格
            style = self.style_combo.currentText()
            
            # 生成AI回复（模拟）
            ai_response = self.generate_ai_response(user_message, style)
            
            # 添加AI回复到历史
            self.add_message_to_history("ai", ai_response)
            
            # 保存到对话历史
            if self.save_history_check.isChecked():
                self.conversation_history.append({
                    "timestamp": datetime.now().isoformat(),
                    "user": user_message,
                    "ai": ai_response,
                    "style": style
                })
            
            # 自动滚动
            if self.auto_scroll_check.isChecked():
                self.scroll_to_bottom()
                
            self.chat_status.setText("回复完成")
            
        except Exception as e:
            self.chat_status.setText(f"处理出错: {str(e)}")
            QMessageBox.critical(self, "错误", f"处理回复时出错: {str(e)}")
        finally:
            # 重新启用发送按钮
            self.send_button.setEnabled(True)
            
    def generate_ai_response(self, message, style):
        """生成AI回复（模拟）"""
        # 根据风格生成不同的回复
        style_responses = {
            "简洁": [
                f"关于'{message}'，我的回答是：这是一个值得思考的问题。",
                f"对于'{message}'，简要来说：需要进一步分析。",
                f"{message}？这个问题可以这样理解。"
            ],
            "详细": [
                f"您提出的问题'{message}'非常有意思。让我详细分析一下：首先，这个问题涉及到多个层面...",
                f"关于'{message}'，这是一个复杂的话题。从历史背景来看...，从技术角度分析...，最后总结...",
                f"对于'{message}'，我将从以下几个方面进行详细阐述：1) 概念定义 2) 实际应用 3) 未来展望"
            ],
            "专业": [
                f"基于当前的研究和实践，对于'{message}'，专业观点是：需要综合考虑多个因素。",
                f"从专业角度分析'{message}'：相关研究表明...，行业标准建议...，最佳实践是...",
                f"关于'{message}'的专业解读：这个问题在学术领域有深入研究，主要结论包括..."
            ],
            "友好": [
                f"嗨！您问的'{message}'真是个有趣的问题！让我来帮您分析一下～",
                f"感谢您的提问！'{message}'这个问题我也很感兴趣，让我分享一下我的看法：",
                f"您提出了一个很好的问题：'{message}'！我觉得可以从这些角度来思考..."
            ],
            "创意": [
                f"想象一下，'{message}'就像夜空中的星星，每一颗都有独特的光芒...",
                f"如果把'{message}'比作一首诗，那么它的韵律是...，意境是...",
                f"关于'{message}'，我有一个创意的视角：让我们换个角度思考这个问题..."
            ]
        }
        
        responses = style_responses.get(style, style_responses["简洁"])
        return random.choice(responses)
        
    def add_message_to_history(self, sender, message):
        """添加消息到对话历史"""
        cursor = self.output_text.textCursor()
        cursor.movePosition(QTextCursor.MoveOperation.End)
        
        # 设置消息样式
        format = QTextCharFormat()
        
        if sender == "user":
            format.setForeground(QColor("#2c3e50"))  # 深蓝色
            format.setFontWeight(QFont.Weight.Bold)
            prefix = "您: "
        else:
            format.setForeground(QColor("#27ae60"))  # 绿色
            format.setFontItalic(True)
            prefix = "DeepSeek: "
            
        # 添加时间戳
        time_str = datetime.now().strftime("%H:%M:%S")
        cursor.insertText(f"[{time_str}] ", format)
        
        # 添加消息
        cursor.insertText(f"{prefix}{message}\n\n", format)
        
    def scroll_to_bottom(self):
        """滚动到底部"""
        scrollbar = self.output_text.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())
        
    def clear_conversation(self):
        """清空对话历史"""
        reply = QMessageBox.question(
            self, "确认清空",
            "确定要清空所有对话历史吗？",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            self.output_text.clear()
            self.conversation_history.clear()
            self.chat_status.setText("对话历史已清空")
            
    def export_conversation(self):
        """导出对话历史为Word文档"""
        if not self.conversation_history:
            QMessageBox.information(self, "无数据", "没有对话历史可以导出")
            return
            
        # 选择保存位置
        from PyQt6.QtWidgets import QFileDialog
        import docx
        from datetime import datetime
        
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "导出对话历史",
            f"DeepSeek对话记录_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            "Word文档 (*.docx)"
        )
        
        if not file_path:
            return
            
        try:
            # 创建Word文档
            doc = docx.Document()
            
            # 添加标题
            doc.add_heading('DeepSeek 对话记录', 0)
            
            # 添加基本信息
            doc.add_paragraph(f'导出时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'对话轮数: {len(self.conversation_history)}')
            doc.add_paragraph('')
            
            # 添加对话内容
            doc.add_heading('对话详情', level=1)
            
            for entry in self.conversation_history:
                # 添加时间戳
                timestamp = datetime.fromisoformat(entry['timestamp']).strftime("%H:%M:%S")
                doc.add_paragraph(f"[{timestamp}] {entry['style']}风格", style='Heading 2')
                
                # 添加用户消息
                doc.add_paragraph(f"用户: {entry['user']}", style='Normal')
                
                # 添加AI回复
                doc.add_paragraph(f"DeepSeek: {entry['ai']}", style='Normal')
                
                # 添加分隔线
                doc.add_paragraph('')
            
            # 保存文档
            doc.save(file_path)
            
            QMessageBox.information(self, "导出成功", f"对话历史已导出至:\n{file_path}")
            self.chat_status.setText("对话历史导出完成")
            
        except Exception as e:
            QMessageBox.critical(self, "导出失败", f"导出过程中发生错误:\n{str(e)}")
            
    def show_conversation_history(self):
        """显示对话历史"""
        if not self.conversation_history:
            QMessageBox.information(self, "无数据", "暂无对话历史")
            return
            
        # 直接调用导出功能
        self.export_conversation()
        
    def closeEvent(self, event):
        """关闭窗口事件"""
        if self.conversation_history and self.save_history_check.isChecked():
            reply = QMessageBox.question(
                self, "保存对话",
                "是否保存当前对话历史？",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No | QMessageBox.StandardButton.Cancel,
                QMessageBox.StandardButton.Yes
            )
            
            if reply == QMessageBox.StandardButton.Cancel:
                event.ignore()
                return
            elif reply == QMessageBox.StandardButton.Yes:
                self.export_conversation()
                
        event.accept()


def main():
    """主函数"""
    app = QApplication(sys.argv)
    app.setApplicationName("DeepSeek Qt浏览器 - 增强版")
    
    # 设置应用程序样式
    app.setStyle("Fusion")
    
    # 设置调色板
    palette = app.palette()
    palette.setColor(QPalette.ColorRole.Window, QColor(240, 240, 240))
    palette.setColor(QPalette.ColorRole.WindowText, QColor(0, 0, 0))
    palette.setColor(QPalette.ColorRole.Base, QColor(255, 255, 255))
    palette.setColor(QPalette.ColorRole.AlternateBase, QColor(245, 245, 245))
    app.setPalette(palette)
    
    window = EnhancedDeepSeekBrowser()
    window.show()
    
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
