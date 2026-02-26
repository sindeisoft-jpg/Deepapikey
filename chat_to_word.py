#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
对话转Word文档工具
"""

import sys
import os
from datetime import datetime
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QVBoxLayout, QHBoxLayout, 
    QWidget, QTextEdit, QPushButton, QLabel, QFileDialog,
    QMessageBox
)
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QFont
import docx
from docx.shared import Inches

class ChatToWordApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.chat_history = []
        self.init_ui()
        
    def init_ui(self):
        self.setWindowTitle('智能对话转Word文档')
        self.setGeometry(100, 100, 800, 600)
        
        # 主窗口部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)
        
        # 标题
        title_label = QLabel('💬 智能对话转Word文档')
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title_font = QFont()
        title_font.setPointSize(16)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)
        
        # 对话显示区域
        self.chat_display = QTextEdit()
        self.chat_display.setReadOnly(True)
        self.chat_display.setStyleSheet("""
            QTextEdit {
                background-color: #f0f0f0;
                border: 1px solid #ccc;
                border-radius: 5px;
                padding: 10px;
                font-family: 'Microsoft YaHei', sans-serif;
                font-size: 12px;
            }
        """)
        layout.addWidget(QLabel('对话记录：'))
        layout.addWidget(self.chat_display)
        
        # 输入区域
        input_layout = QHBoxLayout()
        self.input_field = QTextEdit()
        self.input_field.setMaximumHeight(80)
        self.input_field.setPlaceholderText("请输入您的对话内容...")
        input_layout.addWidget(self.input_field)
        
        send_btn = QPushButton('发送')
        send_btn.clicked.connect(self.send_message)
        send_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        input_layout.addWidget(send_btn)
        layout.addLayout(input_layout)
        
        # 控制按钮
        btn_layout = QHBoxLayout()
        
        clear_btn = QPushButton('清空对话')
        clear_btn.clicked.connect(self.clear_chat)
        clear_btn.setStyleSheet("QPushButton { background-color: #ff9800; color: white; }")
        
        export_btn = QPushButton('导出Word文档')
        export_btn.clicked.connect(self.export_to_word)
        export_btn.setStyleSheet("QPushButton { background-color: #2196F3; color: white; }")
        
        btn_layout.addWidget(clear_btn)
        btn_layout.addWidget(export_btn)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)
        
    def send_message(self):
        message = self.input_field.toPlainText().strip()
        if not message:
            return
            
        # 添加到对话历史
        timestamp = datetime.now().strftime("%H:%M:%S")
        chat_entry = f"[{timestamp}] 用户: {message}"
        self.chat_history.append(chat_entry)
        
        # 显示在界面上
        current_text = self.chat_display.toPlainText()
        if current_text:
            current_text += "\n"
        current_text += chat_entry
        self.chat_display.setPlainText(current_text)
        
        # 清空输入框
        self.input_field.clear()
        
        # 自动滚动到底部
        scrollbar = self.chat_display.verticalScrollBar()
        scrollbar.setValue(scrollbar.maximum())
        
    def clear_chat(self):
        self.chat_display.clear()
        self.chat_history.clear()
        QMessageBox.information(self, "提示", "对话记录已清空")
        
    def export_to_word(self):
        if not self.chat_history:
            QMessageBox.warning(self, "警告", "没有对话内容可导出")
            return
            
        # 选择保存位置
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "保存Word文档",
            f"对话记录_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            "Word文档 (*.docx)"
        )
        
        if not file_path:
            return
            
        try:
            # 创建Word文档
            doc = docx.Document()
            
            # 添加标题
            doc.add_heading('智能对话记录', 0)
            
            # 添加基本信息
            doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'对话条数: {len(self.chat_history)}')
            doc.add_paragraph('')
            
            # 添加对话内容
            doc.add_heading('对话详情', level=1)
            
            for entry in self.chat_history:
                doc.add_paragraph(entry, style='Normal')
            
            # 保存文档
            doc.save(file_path)
            
            QMessageBox.information(self, "成功", f"文档已保存至:\n{file_path}")
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出失败: {str(e)}")

def main():
    app = QApplication(sys.argv)
    
    # 设置应用样式
    app.setStyleSheet("""
        QMainWindow {
            background-color: #ffffff;
        }
        QLabel {
            color: #333333;
            font-family: 'Microsoft YaHei', sans-serif;
        }
    """)
    
    window = ChatToWordApp()
    window.show()
    
    sys.exit(app.exec())

if __name__ == '__main__':
    main()